# core/document_generator.py
#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import copy
import shutil
import logging
from datetime import datetime, timedelta
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree
from .config import Config

logger = logging.getLogger("softdoc_generator")


class DocumentGenerator:
    """文档生成器 - 基于模板复制和文字替换"""
    
    def __init__(self, config: Optional[Config] = None):
        self.config = config or Config()
        self.template_dir = ""
        self.output_dir = ""
    
    def set_template_dir(self, template_dir: str):
        self.template_dir = template_dir
        logger.info(f"设置模板目录: {template_dir}")
    
    def set_output_dir(self, output_dir: str):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        logger.info(f"设置输出目录: {output_dir}")
    
    def generate_documents(self, game_info: Dict, soft_info: Dict, softdoc_files: List[str] = None) -> List[str]:
        """
        生成文档。
        
        Args:
            game_info: 游戏信息字典
            soft_info: 软著信息字典
            softdoc_files: 软著文件列表（支持 PDF 和图片），为空则跳过软著图片输出
        """
        generated_files = []
        
        game_name = self._sanitize_filename(game_info.get('game_name', '未知游戏'))
        game_output_dir = os.path.join(self.output_dir, game_name)
        os.makedirs(game_output_dir, exist_ok=True)
        
        replace_data = self._prepare_replace_data(game_info, soft_info)
        
        templates = [
            ("单机承诺函.docx", "单机承诺函"),
            ("免责承诺函.docx", "免责承诺函"),
        ]
        
        for template_name, doc_type in templates:
            output_file = self._copy_and_replace_template(
                template_name, doc_type, game_output_dir, replace_data, game_info
            )
            if output_file:
                generated_files.append(output_file)
        
        if self._need_authorization(game_info, soft_info):
            auth_file = self._copy_and_replace_template(
                "授权书.docx", "授权书", game_output_dir, replace_data, game_info, soft_info
            )
            if auth_file:
                generated_files.append(auth_file)
        
        # log_file = self._generate_process_log(game_info, soft_info, game_output_dir, generated_files)
        # if log_file:
        #     generated_files.append(log_file)
        
        # 处理软著文件（PDF 转图片，图片直接复制）
        if softdoc_files:
            # 统计已处理的文件数量（用于命名序号）
            file_counter = 1
            for softdoc_path in softdoc_files:
                if not os.path.exists(softdoc_path):
                    logger.warning(f"软著文件不存在，跳过: {softdoc_path}")
                    continue
                
                ext = os.path.splitext(softdoc_path)[1].lower()
                
                if ext == '.pdf':
                    # PDF：转换为图片
                    images = self._convert_pdf_to_images(softdoc_path, game_output_dir, file_counter)
                    generated_files.extend(images)
                    file_counter += len(images)
                else:
                    # 图片：直接复制到输出目录
                    saved = self._copy_softdoc_image(softdoc_path, game_output_dir, file_counter)
                    if saved:
                        generated_files.append(saved)
                        file_counter += 1
        
        logger.info(f"文档生成完成，共生成{len(generated_files)}个文件")
        return generated_files
    
    def _convert_pdf_to_images(self, pdf_path: str, output_dir: str, start_counter: int = 1) -> List[str]:
        """将 PDF 转换为图片并保存到输出目录（纯本地操作，无需 API Key）"""
        try:
            import fitz  # PyMuPDF，纯本地库，不需要 API Key
        except ImportError:
            logger.error("未安装 PyMuPDF，请运行: pip install pymupdf")
            return []
        
        try:
            doc = fitz.open(pdf_path)
            if doc.page_count == 0:
                logger.warning(f"PDF 文件为空: {pdf_path}")
                return []
            
            # 高清渲染：使用 2x 缩放（DPI = 72 * zoom）
            zoom = 2.0
            matrix = fitz.Matrix(zoom, zoom)
            
            saved_images = []
            game_name = os.path.basename(output_dir)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                pix = page.get_pixmap(matrix=matrix)
                
                # 保存为 PNG（支持透明通道，质量更好）
                counter = start_counter + page_num
                new_name = f"{game_name}_软著_{counter}.png"
                new_path = os.path.join(output_dir, new_name)
                pix.save(new_path)
                saved_images.append(new_path)
                
                logger.info(f"保存软著图片: {new_name} (页面 {page_num + 1}/{len(doc)})")
            
            doc.close()
            logger.info(f"PDF 转图片完成，共转换 {len(saved_images)} 页")
            return saved_images
        except Exception as e:
            logger.error(f"PDF 转图片失败: {e}")
            return []
    
    def _copy_softdoc_image(self, image_path: str, output_dir: str, counter: int) -> Optional[str]:
        """复制软著图片到输出目录（保持原格式）"""
        try:
            game_name = os.path.basename(output_dir)
            ext = os.path.splitext(image_path)[1].lower()
            new_name = f"{game_name}_软著_{counter}{ext}"
            new_path = os.path.join(output_dir, new_name)
            
            shutil.copy2(image_path, new_path)
            logger.info(f"保存软著图片: {new_name}")
            return new_path
        except Exception as e:
            logger.error(f"复制软著图片失败: {e}")
            return None
    
    def _prepare_replace_data(self, game_info: Dict, soft_info: Dict) -> Dict:
        auth_years = self.config.get('processing.authorization_years', 10)
        now = datetime.now()
        auth_end_date = now + timedelta(days=auth_years * 365)
        platform_name = self.config.get('processing.platform_name', '广东天宸网络科技有限公司')

        # 著作权人/授权方：去除"企业名称"冗余前缀（支持字间有空格：企 业 名 称:）
        import re as _re
        def _clean_company(name: str) -> str:
            if not name:
                return name
            name = _re.sub(r'^[\u4e00-\u9fff\s]*企[\s]*业[\s]*名[\s]*称[\s]*[：:\s]*', '', name).strip()
            return name

        authorizer = _clean_company(soft_info.get('copyright_holder', game_info.get('developer', '')))
        copyright_holder = _clean_company(soft_info.get('copyright_holder', ''))

        result = {
            '{平台名称}': platform_name,
            '{游戏名称}': game_info.get('game_name', ''),
            '{包名}': game_info.get('package_name', ''),
            '{公司名称}': game_info.get('publisher', ''),
            '{开发者}': game_info.get('developer', ''),
            '{版本}': game_info.get('version', ''),
            '{软件名称}': soft_info.get('software_name', ''),
            '{版本号}': soft_info.get('version', ''),
            '{著作权人}': copyright_holder,
            '{软著登记号}': soft_info.get('registration_number', ''),
            '{授权方}': authorizer,
            '{被授权方}': game_info.get('publisher', ''),
            '{授权年限}': str(auth_years),
            '{授权开始日期}': self._format_date_chinese(now),
            '{授权结束日期}': self._format_date_chinese(auth_end_date),
            '{当前日期}': self._format_date_chinese(now),
            '{当前年份}': str(now.year),
            '{当前月份}': str(now.month),
            '{当前日}': str(now.day),
            '{版号}': self.config.get('processing.game_version_number', '0'),
        }
        
        # 调试输出
        print("\n=== 替换数据 ===")
        for key, value in result.items():
            if '当前' in key or '年份' in key or '月份' in key or '日' in key:
                print(f"  {key} -> {value}")
        
        return result

    # 需要关闭 Word 拼写检查的占位符（通常是包名、版本号等英文内容）
    NO_PROOF_KEYS = {'{包名}', '{版本}', '{版本号}', '{软著登记号}'}

    
    def _format_date_chinese(self, date: datetime) -> str:
        return f"{date.year}年{date.month}月{date.day}日"
    
    def _copy_and_replace_template(self, template_name: str, doc_type: str,
                                    output_dir: str, replace_data: Dict,
                                    game_info: Dict, soft_info: Dict = None) -> Optional[str]:
        template_path = os.path.join(self.template_dir, template_name)
        
        if not os.path.exists(template_path):
            logger.warning(f"模板不存在: {template_path}，跳过 {doc_type}")
            return None
        
        game_name = self._sanitize_filename(game_info.get('game_name', '游戏'))
        company = self._sanitize_filename(game_info.get('publisher', '公司'))
        
        if doc_type == "单机承诺函":
            filename = f"{game_name}-单机-{company}.docx"
        elif doc_type == "免责承诺函":
            filename = f"{game_name}-免责-{company}.docx"
        elif doc_type == "授权书":
            licensor = self._sanitize_filename((soft_info or {}).get('copyright_holder', '授权方'))
            filename = f"{game_name}-授权书-{company}-{licensor}.docx"
        else:
            filename = f"{game_name}-{doc_type}-{company}.docx"
        
        filename = self._sanitize_filename(filename)
        output_path = os.path.join(output_dir, filename)
        
        try:
            shutil.copy2(template_path, output_path)
            logger.info(f"复制模板: {template_name}")
            
            doc = Document(output_path)
            replace_count = self._replace_all_text(doc, replace_data)
            doc.save(output_path)
            
            logger.info(f"完成 {doc_type}: 替换了 {replace_count} 处，输出: {filename}")
            return output_path
        except Exception as e:
            logger.error(f"处理模板 {template_name} 失败: {e}")
            return None
    
    def _replace_all_text(self, doc, replace_data: Dict) -> int:
        replace_count = 0

        for paragraph in doc.paragraphs:
            replace_count += self._replace_in_paragraph(paragraph, replace_data)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_count += self._replace_in_paragraph(paragraph, replace_data)

        if hasattr(doc, 'sections'):
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        replace_count += self._replace_in_paragraph(paragraph, replace_data)
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        replace_count += self._replace_in_paragraph(paragraph, replace_data)

        return replace_count

    # ------------------------------------------------------------------ #
    # 以下三个旧方法保留但不再使用，避免破坏其他调用链
    # ------------------------------------------------------------------ #
    def _get_run_format(self, run) -> Dict:
        """获取 run 的完整格式信息（保留兼容）"""
        format_info = {
            'font_name': None, 'font_size': None,
            'bold': False, 'italic': False, 'underline': False,
            'strike': False, 'color': None, 'highlight': None,
        }
        try: format_info['font_name'] = run.font.name
        except: pass
        try:
            if run.font.size: format_info['font_size'] = run.font.size
        except: pass
        try: format_info['bold'] = run.bold if run.bold is not None else False
        except: pass
        try: format_info['italic'] = run.italic if run.italic is not None else False
        except: pass
        try: format_info['underline'] = run.font.underline if run.font.underline is not None else False
        except: pass
        try: format_info['strike'] = run.font.strike if hasattr(run.font, 'strike') and run.font.strike is not None else False
        except: pass
        try:
            if run.font.color and run.font.color.rgb:
                format_info['color'] = run.font.color.rgb
        except: pass
        try: format_info['highlight'] = run.font.highlight_color
        except: pass
        return format_info

    def _apply_run_format(self, run, format_info: Dict):
        """应用格式到 run（保留兼容）"""
        if format_info['font_name']:
            try: run.font.name = format_info['font_name']
            except: pass
        if format_info['font_size']:
            try: run.font.size = format_info['font_size']
            except: pass
        try: run.bold = format_info['bold']
        except: pass
        try: run.italic = format_info['italic']
        except: pass
        try: run.font.underline = format_info['underline']
        except: pass
        try:
            if format_info['strike']: run.font.strike = format_info['strike']
        except: pass
        if format_info['color']:
            try: run.font.color.rgb = format_info['color']
            except: pass
        if format_info['highlight']:
            try: run.font.highlight_color = format_info['highlight']
            except: pass

    # ------------------------------------------------------------------ #
    # 核心：基于 XML 节点的段落文本替换（保留所有格式）
    # ------------------------------------------------------------------ #
    def _replace_in_paragraph(self, paragraph, replace_data: Dict) -> int:
        """
        使用 XML 级别操作替换段落文本，完整保留下划线、加粗、斜体等格式。

        策略：
        1. 先判断段落完整文本中是否含有占位符，没有则快速返回。
        2. 收集所有 <w:r> 和其文本，拼出全文，在全文中定位占位符的字符范围。
        3. 将字符范围映射回具体的 run 区间，对这组 run 进行合并（只保留首个
           run 的 <w:rPr>），然后只修改 <w:t> 文本，<w:rPr> 完全不动。
        4. 不跨 run 的占位符直接在原 run 上原地替换文本。
        5. 对 NO_PROOF_KEYS 中的占位符，替换后自动给 run 加 <w:noProof/>，
           避免 Word 对包名等英文内容显示红色拼写波浪线。
        """
        full_text = paragraph.text
        if not full_text:
            return 0
        need_replace = any(k in full_text for k in replace_data)
        if not need_replace:
            return 0

        W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        p_elem = paragraph._p

        # 收集 <w:r> 元素及每个 run 的文本
        def collect_runs():
            rs = p_elem.findall(f'{{{W}}}r')
            runs = []
            for r in rs:
                txt = ''.join(t.text or '' for t in r.findall(f'{{{W}}}t'))
                runs.append({'elem': r, 'text': txt})
            return runs

        replace_count = 0

        # 迭代替换，直到没有更多占位符（每次替换后 run 列表可能改变）
        max_iter = 20
        for _ in range(max_iter):
            runs = collect_runs()
            if not runs:
                break

            # 构建全文及每个字符所属的 run 索引
            char_to_run = []  # char_to_run[i] = run index
            full = ''
            for ri, r in enumerate(runs):
                full += r['text']
                char_to_run.extend([ri] * len(r['text']))

            # 在全文中找第一个占位符
            found_old = None
            found_start = None
            for old in replace_data:
                pos = full.find(old)
                if pos != -1:
                    if found_start is None or pos < found_start:
                        found_start = pos
                        found_old = old

            if found_old is None:
                break  # 没有更多占位符了

            found_end = found_start + len(found_old) - 1  # 包含的最后一个字符索引
            new_val = replace_data[found_old]
            replace_count += 1
            need_no_proof = found_old in self.NO_PROOF_KEYS

            ri_start = char_to_run[found_start]
            ri_end = char_to_run[found_end]

            if ri_start == ri_end:
                # 占位符在同一个 run 内，直接替换文本
                r_elem = runs[ri_start]['elem']
                old_text = runs[ri_start]['text']
                new_text = old_text.replace(found_old, new_val, 1)
                self._set_run_text(r_elem, new_text, W)
                if need_no_proof:
                    self._ensure_no_proof(r_elem, W)
            else:
                # 占位符跨多个 run，需要合并
                merged_text = ''.join(runs[k]['text'] for k in range(ri_start, ri_end + 1))
                new_merged = merged_text.replace(found_old, new_val, 1)

                # 修改 ri_start 的文本
                r_elem = runs[ri_start]['elem']
                self._set_run_text(r_elem, new_merged, W)
                if need_no_proof:
                    self._ensure_no_proof(r_elem, W)

                # 删除 ri_start+1 .. ri_end 的 run
                for k in range(ri_end, ri_start, -1):
                    p_elem.remove(runs[k]['elem'])

        return replace_count

    def _ensure_no_proof(self, r_elem, W: str):
        """给 run 的 <w:rPr> 加上 <w:noProof/>，关闭 Word 拼写/语法检查"""
        rpr = r_elem.find(f'{{{W}}}rPr')
        if rpr is None:
            # 没有 <w:rPr>，创建一个并插到 run 最前面
            rpr = etree.Element(f'{{{W}}}rPr')
            r_elem.insert(0, rpr)
        # 检查是否已有 <w:noProof>
        no_proof = rpr.find(f'{{{W}}}noProof')
        if no_proof is None:
            etree.SubElement(rpr, f'{{{W}}}noProof')


    def _merge_runs_and_replace(self, p_elem, replace_data: Dict, W: str) -> int:
        """兼容保留，实际逻辑已移入 _replace_in_paragraph"""
        return 0

    def _set_run_text(self, r_elem, text: str, W: str):
        """将 run 的所有 <w:t> 合并为一个，内容设为 text，不修改 <w:rPr>"""
        t_elems = r_elem.findall(f'{{{W}}}t')
        if t_elems:
            t_elems[0].text = text
            if text != text.strip():
                t_elems[0].set(
                    '{http://www.w3.org/XML/1998/namespace}space', 'preserve'
                )
            for t in t_elems[1:]:
                r_elem.remove(t)
        else:
            t_new = etree.SubElement(r_elem, f'{{{W}}}t')
            t_new.text = text
            if text != text.strip():
                t_new.set(
                    '{http://www.w3.org/XML/1998/namespace}space', 'preserve'
                )

    # 保留旧方法名供外部兼容
    def _replace_text_in_paragraph(self, paragraph, replace_data: Dict) -> int:
        return self._replace_in_paragraph(paragraph, replace_data)


    
    def _need_authorization(self, game_info: Dict, soft_info: Dict) -> bool:
        publisher = game_info.get('publisher', '')
        copyright_holder = soft_info.get('copyright_holder', '')
        
        if publisher and copyright_holder and publisher != copyright_holder:
            logger.info(f"上架主体与著作权人不一致，需要生成授权书")
            return True
        
        return False
    
    def _generate_process_log(self, game_info: Dict, soft_info: Dict, output_dir: str, generated_files: List[str]) -> Optional[str]:
        try:
            log_path = os.path.join(output_dir, "处理日志.txt")
            
            with open(log_path, 'w', encoding='utf-8-sig') as f:
                f.write("=" * 60 + "\n")
                f.write("软著文档生成工具 - 处理日志\n")
                f.write("=" * 60 + "\n\n")
                f.write(f"处理时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                f.write("游戏信息:\n")
                f.write(f"  游戏名称: {game_info.get('game_name', '')}\n")
                f.write(f"  包名: {game_info.get('package_name', '')}\n")
                f.write(f"  上架主体: {game_info.get('publisher', '')}\n\n")
                f.write("软著信息:\n")
                f.write(f"  软件名称: {soft_info.get('software_name', '')}\n")
                f.write(f"  著作权人: {soft_info.get('copyright_holder', '')}\n")
                f.write(f"  登记号: {soft_info.get('registration_number', '')}\n\n")
                f.write("生成的文件:\n")
                for i, file_path in enumerate(generated_files, 1):
                    f.write(f"  {i}. {os.path.basename(file_path)}\n")
                f.write("\n" + "=" * 60 + "\n")
            
            return log_path
        except Exception as e:
            logger.error(f"生成日志失败: {e}")
            return None
    
    def _sanitize_filename(self, filename: str) -> str:
        if not filename:
            return "未命名"
        
        illegal_chars = r'<>:"/\|?*'
        for char in illegal_chars:
            filename = filename.replace(char, '_')
        
        filename = filename.strip()
        
        if len(filename) > 200:
            filename = filename[:200]
        
        if not filename:
            filename = "未命名"
        
        return filename


def main():
    logging.basicConfig(level=logging.INFO)
    
    game_info = {
        'game_name': '生存挑战模拟',
        'package_name': 'com.hxwl.sctzmn.vivominigame',
        'publisher': '深圳市鸿鑫网络科技有限公司',
        'developer': '深圳市顺思畅想科技有限公司',
        'version': '1.0.0',
    }
    
    soft_info = {
        'software_name': '生存挑战模拟游戏软件',
        'version': 'V1.0',
        'copyright_holder': '深圳市顺思畅想科技有限公司',
        'registration_number': '2021SRE028650',
    }
    
    generator = DocumentGenerator()
    generator.set_template_dir("./templates")
    generator.set_output_dir("./test_output")
    files = generator.generate_documents(game_info, soft_info)
    
    print(f"生成了 {len(files)} 个文件")


if __name__ == "__main__":
    main()